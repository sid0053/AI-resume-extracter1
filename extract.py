import argparse
import os
import re
from typing import Tuple, List

# PDF
import pdfplumber

# DOCX
from docx import Document

# ---------- Cleaning ----------

# Keep common tech punctuation: + . - # / _ & @
_ALLOWED_CHARS = r"[^a-z0-9@\.\+\-#/_&\s\n]"
BULLETS = ["•", "◦", "·", "●", "–", "—", "•", "▪", "‣"]

def normalize_bullets(text: str) -> str:
    for b in BULLETS:
        text = text.replace(b, "\n- ")
    return text

def de_duplicate_headers_footers(lines: List[str]) -> List[str]:
    """
    Heuristic: if a line appears on > 40% of pages at same position,
    treat it as header/footer and drop it.
    Works best when called per page and then across all pages.
    Here we apply a simple global de-dup across the whole doc.
    """
    freq = {}
    for ln in lines:
        k = ln.strip()
        if not k:
            continue
        freq[k] = freq.get(k, 0) + 1
    if not freq:
        return lines
    threshold = max(2, int(0.4 * max(freq.values())))
    return [ln for ln in lines if freq.get(ln.strip(), 0) < threshold]

def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = normalize_bullets(text)
    text = text.lower()
    text = re.sub(_ALLOWED_CHARS, " ", text)       # strip disallowed chars
    text = re.sub(r"[ \t]{2,}", " ", text)         # collapse spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)         # limit blank lines
    # keep hyphenated tech tokens as-is; ensure space around slashes when needed
    return text.strip()

# ---------- PDF ----------

def _page_text_safe(page) -> str:
    txt = page.extract_text(x_tolerance=2, y_tolerance=2)
    return txt or ""

def is_scanned_pdf(pdf_path: str, sample_pages: int = 3) -> bool:
    try:
        with pdfplumber.open(pdf_path) as pdf:
            pages = pdf.pages[:sample_pages] if len(pdf.pages) > sample_pages else pdf.pages
            empty_count = 0
            for p in pages:
                if not _page_text_safe(p).strip():
                    empty_count += 1
            # mostly empty text across samples -> likely scanned
            return empty_count >= max(1, len(pages) - 1)
    except Exception:
        return False

def extract_text_from_pdf(path: str) -> str:
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = _page_text_safe(page)
            if not t:
                continue
            # split to lines to help downstream header/footer trimming
            texts.append(t)
    out = "\n".join(texts)
    # optional: mild de-dup of repeated lines
    lines = out.splitlines()
    lines = de_duplicate_headers_footers(lines)
    return "\n".join(lines)

# ---------- DOCX ----------

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = []

    # paragraphs
    for para in doc.paragraphs:
        parts.append(para.text)

    # tables (often list skills/experience)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)

# ---------- Public API ----------

def extract_resume_text(file_path: str, clean: bool = True) -> Tuple[str, str]:
    """
    Returns (raw_text, cleaned_text)
    cleaned_text may equal raw_text if clean=False.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        raw = extract_text_from_pdf(file_path)
    elif ext in (".docx",):
        raw = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported format. Use PDF or DOCX.")

    if clean:
        return raw, clean_text(raw)
    else:
        return raw, raw

# ---------- CLI ----------

def main():
    parser = argparse.ArgumentParser(description="Extract text from a resume (PDF/DOCX) and save cleaned text.")
    parser.add_argument("input", help="Path to resume (.pdf or .docx)")
    parser.add_argument("-o", "--output", default="extracted_resume.txt", help="Output .txt path (default: extracted_resume.txt)")
    parser.add_argument("--no-clean", action="store_true", help="Do not clean text (save raw)")
    args = parser.parse_args()

    try:
        if os.path.splitext(args.input)[1].lower() == ".pdf" and is_scanned_pdf(args.input):
            print("⚠️  This looks like a scanned PDF (image-only). Text may be empty without OCR.")

        raw, cleaned = extract_resume_text(args.input, clean=not args.no_clean)
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(cleaned)

        print(f"✅ Extracted {'cleaned' if not args.no_clean else 'raw'} text saved to {args.output}")
        if not cleaned.strip():
            print("⚠️  Output is empty. If it's a scanned PDF, consider OCR (pytesseract + pdf2image).")

    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()
